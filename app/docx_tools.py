from __future__ import annotations

import json
from collections import Counter
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document


def extract_text(data: bytes) -> str:
    document = Document(BytesIO(data))
    blocks = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append("\t".join(cells))
    return "\n".join(blocks)


def _cm(value: Any) -> float | None:
    return round(value.cm, 2) if value is not None else None


def _finding(
    findings: list[dict],
    item: str,
    expected: Any,
    actual: Any,
    status: str,
    evidence: str,
) -> None:
    findings.append({
        "item": item,
        "expected": expected,
        "actual": actual,
        "status": status,
        "evidence": evidence,
    })


def _compare_number(actual: float | None, expected: float, tolerance: float) -> str:
    if actual is None:
        return "unknown"
    return "pass" if abs(actual - expected) <= tolerance else "fail"


def _font_name(run) -> str | None:
    if run.font.name:
        return run.font.name
    rpr = run._element.rPr
    if rpr is not None and rpr.rFonts is not None:
        return rpr.rFonts.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia")
    return None


def _paragraph_style(paragraph) -> dict:
    runs = [run for run in paragraph.runs if run.text.strip()]
    names = [_font_name(run) for run in runs if _font_name(run)]
    sizes = [round(run.font.size.pt, 1) for run in runs if run.font.size]
    bolds = [run.bold for run in runs if run.bold is not None]
    return {
        "font": Counter(names).most_common(1)[0][0] if names else None,
        "size_pt": Counter(sizes).most_common(1)[0][0] if sizes else None,
        "bold": Counter(bolds).most_common(1)[0][0] if bolds else None,
        "first_line_indent_cm": _cm(paragraph.paragraph_format.first_line_indent),
        "line_spacing": paragraph.paragraph_format.line_spacing,
    }


def inspect_document(data: bytes, policy_path: Path) -> dict:
    policy = json.loads(policy_path.read_text(encoding="utf-8"))
    document = Document(BytesIO(data))
    paragraphs = [p for p in document.paragraphs if p.text.strip()]
    text = "\n".join(p.text.strip() for p in paragraphs)
    findings: list[dict] = []

    section = document.sections[0]
    page = policy["page"]
    page_values = {
        "纸张宽度": _cm(section.page_width),
        "纸张高度": _cm(section.page_height),
        "上页边距": _cm(section.top_margin),
        "下页边距": _cm(section.bottom_margin),
        "左页边距": _cm(section.left_margin),
        "右页边距": _cm(section.right_margin),
    }
    expected_values = {
        "纸张宽度": page["width_cm"],
        "纸张高度": page["height_cm"],
        "上页边距": page["top_margin_cm"],
        "下页边距": page["bottom_margin_cm"],
        "左页边距": page["left_margin_cm"],
        "右页边距": page["right_margin_cm"],
    }
    for item, actual in page_values.items():
        expected = expected_values[item]
        _finding(
            findings, item, expected, actual,
            _compare_number(actual, expected, page["tolerance_cm"]),
            "文档第一个节的页面设置",
        )

    positions = []
    for name in policy["required_sections"]:
        position = text.lower().find(name.lower())
        positions.append(position)
        _finding(
            findings,
            f"章节：{name}",
            "存在",
            "存在" if position >= 0 else "缺失",
            "pass" if position >= 0 else "fail",
            f"全文检索位置 {position}",
        )
    existing = [value for value in positions if value >= 0]
    order_ok = len(existing) == len(positions) and existing == sorted(existing)
    _finding(
        findings, "章节顺序", policy["required_sections"],
        "按要求排列" if order_ok else "存在缺失或顺序错误",
        "pass" if order_ok else "fail", "按全文位置比较",
    )

    if paragraphs:
        actual = _paragraph_style(paragraphs[0])
        expected = policy["typography"]["title"]
        title_ok = (
            actual["font"] == expected["font"]
            and actual["size_pt"] == expected["size_pt"]
            and actual["bold"] is expected["bold"]
        )
        known = all(actual[key] is not None for key in ("font", "size_pt", "bold"))
        _finding(
            findings, "论文题目样式", expected, actual,
            "pass" if title_ok else "fail" if known else "unknown",
            f"首个非空段落：{paragraphs[0].text[:40]}",
        )

    body_candidates = [
        p for p in paragraphs
        if len(p.text.strip()) >= 30 and not p.style.name.lower().startswith("heading")
    ]
    if body_candidates:
        actual = _paragraph_style(body_candidates[0])
        expected = policy["typography"]["body"]
        known = actual["font"] is not None and actual["size_pt"] is not None
        font_ok = actual["font"] == expected["font"]
        size_ok = actual["size_pt"] == expected["size_pt"]
        _finding(
            findings, "正文样式", expected, actual,
            "pass" if known and font_ok and size_ok else "fail" if known else "unknown",
            f"正文样本：{body_candidates[0].text[:40]}",
        )
    else:
        _finding(findings, "正文样式", policy["typography"]["body"], None,
                 "unknown", "没有长度达到 30 字符的正文段落")

    counts = Counter(item["status"] for item in findings)
    return {
        "policy_id": policy["policy_id"],
        "summary": {
            "total": len(findings),
            "pass": counts["pass"],
            "fail": counts["fail"],
            "unknown": counts["unknown"],
        },
        "findings": findings,
    }