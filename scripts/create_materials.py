from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parent.parent
RESOURCE_DIR = ROOT / "resources"
RESOURCE_DIR.mkdir(exist_ok=True)


def set_run(run, font: str, size: float, bold: bool = False) -> None:
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold


def create_specification() -> None:
    document = Document()
    title = document.add_paragraph()
    set_run(title.add_run("计算机学院本科毕业论文格式规范"), "黑体", 18, True)
    sections = {
        "一、页面设置": [
            "使用 A4 纸张，宽 21.0 cm，高 29.7 cm。",
            "页边距为上 2.5 cm、下 2.5 cm、左 3.0 cm、右 2.5 cm。",
        ],
        "二、字体与段落": [
            "论文题目使用黑体 18 pt 并加粗。",
            "一级标题使用黑体 16 pt 并加粗。",
            "正文使用宋体 12 pt，首行缩进约 0.74 cm，1.5 倍行距。",
        ],
        "三、章节顺序": [
            "依次设置摘要、关键词、ABSTRACT、目录、绪论、系统设计、结论、参考文献、致谢。"
        ],
        "四、参考文献": [
            "保留作者、题名、来源、年份和 DOI 等可复核字段。",
            "近五年文献比例不得低于 50%。",
        ],
    }
    for heading, rows in sections.items():
        paragraph = document.add_paragraph()
        set_run(paragraph.add_run(heading), "黑体", 16, True)
        for row in rows:
            paragraph = document.add_paragraph(row)
            for run in paragraph.runs:
                set_run(run, "宋体", 12)
    document.save(RESOURCE_DIR / "计算机学院本科毕业论文格式规范.docx")


def create_paper() -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

    title_text = "面向高校课程的多智能体辅助系统设计与实现"
    title = document.add_paragraph()
    set_run(title.add_run(title_text), "宋体", 16, True)
    document.add_paragraph("作者姓名：张晨")
    document.add_paragraph("作者编号：2026230417")
    document.add_paragraph(f"论文题目：{title_text}")

    content = [
        ("摘要", "本文围绕多智能体协作、课程资源检索和过程评价开展系统设计，并完成主要功能验证。"),
        ("关键词", "多智能体；工具调用；知识检索；过程评价"),
        ("ABSTRACT", "This paper designs a multi-agent support system for course activities."),
        ("目录", "1 绪论\n2 系统设计\n3 结论"),
        ("绪论", "相关应用需要把模型决策、外部工具和业务规则组合为可追踪的处理流程。"),
        ("系统设计", "系统由任务编排、工具调用、知识检索和结果展示四个模块组成，各模块通过明确接口协作。"),
        ("结论", "系统完成了核心流程验证，后续可继续补充评价指标和权限控制。"),
    ]
    for heading, body in content:
        paragraph = document.add_paragraph()
        set_run(paragraph.add_run(heading), "黑体", 16, True)
        paragraph = document.add_paragraph(body)
        paragraph.paragraph_format.first_line_indent = Cm(0.74)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        for run in paragraph.runs:
            set_run(run, "宋体", 12)

    heading = document.add_paragraph()
    set_run(heading.add_run("参考文献"), "黑体", 16, True)
    references = [
        "[1] Vaswani A, et al. Attention Is All You Need. 2017.",
        "[2] Brown T B, et al. Language Models are Few-Shot Learners. 2020.",
        "[3] Lewis P, et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks. 2020.",
        "[4] Yao S, et al. ReAct: Synergizing Reasoning and Acting in Language Models. 2023.",
        "[5] Park J S, et al. Generative Agents: Interactive Simulacra of Human Behavior. 2023. DOI:10.1145/3586183.3606763.",
        "[6] Wu Q, et al. AutoGen: Enabling Next-Gen LLM Applications. 2024.",
    ]
    for value in references:
        paragraph = document.add_paragraph(value)
        for run in paragraph.runs:
            set_run(run, "宋体", 12)

    heading = document.add_paragraph()
    set_run(heading.add_run("致谢"), "黑体", 16, True)
    document.add_paragraph("感谢在资料整理和系统测试过程中提供的支持。")
    document.save(RESOURCE_DIR / "论文初稿.docx")


create_specification()
create_paper()
print("已生成格式规范和论文初稿")